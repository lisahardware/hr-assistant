"""
tools/vector_store.py
Manages the ChromaDB vector store for HR templates and compliance documents.
Runs entirely locally — no data leaves the machine.
"""

import os
import json
import hashlib
import chromadb
from chromadb.utils import embedding_functions
from docx import Document


class VectorStore:
    def __init__(self, persist_path: str = None):
        # Default resolution order:
        #   1. explicit persist_path argument, if given
        #   2. CHROMA_PERSIST_PATH environment variable
        #   3. "./hr_vector_store" (original default, for local/non-Docker use)
        #
        # This matters because docker-compose.yml mounts the named
        # "chroma-data" volume at /app/chroma_db specifically so vector
        # store data survives container rebuilds. The original hardcoded
        # default of "./hr_vector_store" resolves (relative to the
        # container's WORKDIR, /app) to /app/hr_vector_store — a
        # completely different, non-persisted path — so all indexed
        # templates and compliance sources were silently being written
        # outside the volume the whole time. Set CHROMA_PERSIST_PATH=
        # /app/chroma_db in the hr-app service's environment (matching
        # the compose file's mount target) to fix this for Docker
        # deployments; the bare relative-path fallback still works
        # correctly for running main.py directly on a host machine.
        if persist_path is None:
            persist_path = os.environ.get("CHROMA_PERSIST_PATH", "./hr_vector_store")

        self.client = chromadb.PersistentClient(path=persist_path)
        
        # Use the default sentence transformer for local embeddings
        # This runs entirely on your machine with no API calls
        self.embedding_fn = embedding_functions.DefaultEmbeddingFunction()

        # Two collections: internal templates and compliance sources
        self.templates = self.client.get_or_create_collection(
            name="hr_templates",
            embedding_function=self.embedding_fn,
            metadata={"description": "Owner HR document templates"}
        )
        self.compliance = self.client.get_or_create_collection(
            name="compliance_sources",
            embedding_function=self.embedding_fn,
            metadata={"description": "Public labor law and compliance references"}
        )

    def _extract_template_text(self, file_path: str) -> str:
        """
        Read a template file's text content, dispatching based on file
        extension. .docx files are zip archives, not plain text, so they
        need python-docx to extract readable content; .txt files are
        read directly as plain text, unchanged from the original
        behavior.
        """
        if file_path.lower().endswith(".docx"):
            doc = Document(file_path)
            lines = [p.text for p in doc.paragraphs]
            # Include table cell text too, since several of this
            # project's templates (contractor agreement, onboarding
            # checklist) put real content inside tables, not just body
            # paragraphs.
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            lines.append(p.text)
            return "\n".join(lines)
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()

    def index_template(self, file_path: str) -> int:
        """
        Index an HR template file into the vector store.
        Chunks by section (lines starting with ##) to preserve
        semantic boundaries rather than splitting arbitrarily.
        Returns the number of chunks indexed.

        Supports both .txt files (read as plain text) and .docx files
        (text extracted via python-docx). A .docx file is a zip archive
        under the hood, not plain text, so opening it with
        open(path, "r", encoding="utf-8") reliably fails with a UTF-8
        decode error on the binary zip bytes — that's what this method
        used to do unconditionally, even though index_all_templates()
        has always looked for both .txt and .docx files.
        """
        content = self._extract_template_text(file_path)

        chunks = self._chunk_by_section(content)
        file_name = os.path.basename(file_path)

        for i, chunk in enumerate(chunks):
            chunk_id = f"{file_name}_chunk_{i}"
            # Check if already indexed to avoid duplicates on re-run
            existing = self.templates.get(ids=[chunk_id])
            if existing["ids"]:
                self.templates.update(
                    ids=[chunk_id],
                    documents=[chunk["content"]],
                    metadatas=[{
                        "source": file_name,
                        "section": chunk["heading"],
                        "file_path": file_path
                    }]
                )
            else:
                self.templates.add(
                    ids=[chunk_id],
                    documents=[chunk["content"]],
                    metadatas=[{
                        "source": file_name,
                        "section": chunk["heading"],
                        "file_path": file_path
                    }]
                )

        return len(chunks)

    def index_compliance_text(self, text: str, source_url: str, state: str, retrieved_date: str) -> int:
        """
        Index a compliance document from a public source.
        Called by the compliance refresh tool on a schedule.
        """
        chunks = self._chunk_by_section(text)

        # Stable, short, unique-per-URL identifier component. Using the
        # source URL itself (or chunk index alone) directly in the ID
        # previously caused a real bug: chunk_id was built as
        # "compliance_{state}_{i}" with no source-specific component, so
        # every source fetched for the same state collided on the same
        # small set of IDs (since each source's chunk-index i restarts
        # at 0). With most sources reducing to a single chunk (no "##"
        # headings in fetched HTML-stripped text), each new source for a
        # state silently overwrote the previous one — a full refresh_all()
        # run would leave only the LAST source's content indexed per
        # state, even though every fetch appeared to succeed in the logs.
        url_hash = hashlib.sha1(source_url.encode("utf-8")).hexdigest()[:10]

        for i, chunk in enumerate(chunks):
            chunk_id = f"compliance_{state}_{url_hash}_{i}"
            existing = self.compliance.get(ids=[chunk_id])
            if existing["ids"]:
                self.compliance.update(
                    ids=[chunk_id],
                    documents=[chunk["content"]],
                    metadatas=[{
                        "source": source_url,
                        "state": state,
                        "retrieved_date": retrieved_date,
                        "section": chunk["heading"]
                    }]
                )
            else:
                self.compliance.add(
                    ids=[chunk_id],
                    documents=[chunk["content"]],
                    metadatas=[{
                        "source": source_url,
                        "state": state,
                        "retrieved_date": retrieved_date,
                        "section": chunk["heading"]
                    }]
                )

        return len(chunks)

    def retrieve_templates(self, query: str, n_results: int = 4, where: dict = None) -> list[dict]:
        """
        Semantic search over internal HR templates.
        Optionally filter by metadata using a where clause.
        """
        count = self.templates.count()
        if count == 0:
            return []
        
        kwargs = {
            "query_texts": [query],
            "n_results": min(n_results, count)
        }
        if where:
            kwargs["where"] = where

        results = self.templates.query(**kwargs)
        return self._format_results(results)

    def retrieve_compliance(self, query: str, state: str = None, n_results: int = 4) -> list[dict]:
        """
        Semantic search over compliance documents.
        Optionally filter by state for jurisdiction-specific results.
        """
        where = {"state": state} if state else None
        count = self.compliance.count()
        if count == 0:
            return []

        results = self.compliance.query(
            query_texts=[query],
            n_results=min(n_results, count),
            where=where
        )
        return self._format_results(results)

    def _chunk_by_section(self, text: str) -> list[dict]:
        """
        Split document text at section headings (## Header).
        Falls back to the full document as one chunk if no headings found.
        """
        lines = text.split("\n")
        chunks = []
        current_heading = "Introduction"
        current_lines = []

        for line in lines:
            if line.startswith("## "):
                if current_lines:
                    chunks.append({
                        "heading": current_heading,
                        "content": "\n".join(current_lines).strip()
                    })
                current_heading = line.replace("## ", "").strip()
                current_lines = []
            else:
                current_lines.append(line)

        if current_lines:
            chunks.append({
                "heading": current_heading,
                "content": "\n".join(current_lines).strip()
            })

        # Filter out empty chunks
        return [c for c in chunks if c["content"]]

    def _format_results(self, results: dict) -> list[dict]:
        """Format ChromaDB query results into a clean list of dicts."""
        formatted = []
        if not results["documents"] or not results["documents"][0]:
            return formatted
        for doc, meta in zip(results["documents"][0], results["metadatas"][0]):
            formatted.append({
                "content": doc,
                "metadata": meta
            })
        return formatted

    def get_compliance_last_refreshed(self) -> str | None:
        """
        Return the most recent retrieved_date across all indexed compliance
        chunks, or None if nothing has been indexed yet. Each chunk stores
        its own retrieved_date in metadata (set by compliance_refresh.py),
        so "last refreshed" is derived as the max of those values rather
        than tracked as a separate piece of state — this stays correct
        automatically even if individual states are refreshed at
        different times via `--state`, since it reflects the most recent
        refresh across all currently-indexed content.
        """
        count = self.compliance.count()
        if count == 0:
            return None

        # Pull just the metadata for every chunk — cheap relative to a
        # full query, since we only need the retrieved_date field, not
        # document content or embeddings.
        all_items = self.compliance.get(include=["metadatas"])
        dates = [
            m.get("retrieved_date")
            for m in all_items.get("metadatas", [])
            if m and m.get("retrieved_date")
        ]
        return max(dates) if dates else None

    def index_all_templates(self, templates_dir: str = None) -> dict:
        """Index all template files in the templates directory."""
        if templates_dir is None:
            base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            templates_dir = os.path.join(base_dir, "templates")
        
        results = {}
        for filename in os.listdir(templates_dir):
            if filename.endswith(".txt") or filename.endswith(".docx"):
                path = os.path.join(templates_dir, filename)
                count = self.index_template(path)
                results[filename] = count
        return results