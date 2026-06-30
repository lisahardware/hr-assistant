"""
tools/document_creator.py
Generates .docx HR documents by filling placeholders in templates.
Uses python-docx to produce real Word files that open in Microsoft Word
and can be saved to OneDrive via the MCP server.
"""

import os
import re
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocumentCreator:
    def __init__(self, templates_dir: str = None, output_dir: str = None):
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        self.templates_dir = templates_dir or os.path.join(base_dir, "templates")
        self.output_dir = output_dir or os.path.join(base_dir, "output")
        os.makedirs(self.output_dir, exist_ok=True)

    def create_offer_letter(self, fields: dict) -> str:
        template_path = os.path.join(self.templates_dir, "offer_letter_template.docx")
        doc = self._fill_template(template_path, fields)
        filename = f"{fields.get('candidate_name', 'candidate').replace(' ', '_')}_offer_letter_{date.today()}.docx"
        return self._save_docx(doc, filename)

    def create_contractor_agreement(self, fields: dict) -> str:
        template_path = os.path.join(self.templates_dir, "contractor_agreement_template.docx")
        doc = self._fill_template(template_path, fields)
        filename = f"{fields.get('contractor_name', 'contractor').replace(' ', '_')}_contractor_agreement_{date.today()}.docx"
        return self._save_docx(doc, filename)

    def create_onboarding_checklist(self, fields: dict) -> str:
        template_path = os.path.join(self.templates_dir, "onboarding_checklist_template.docx")
        doc = self._fill_template(template_path, fields)
        filename = f"{fields.get('candidate_name', 'newhire').replace(' ', '_')}_onboarding_checklist_{date.today()}.docx"
        return self._save_docx(doc, filename)

    def create_job_description(self, fields: dict) -> str:
        template_path = os.path.join(self.templates_dir, "job_description_template.docx")
        doc = self._fill_template(template_path, fields)
        filename = f"{fields.get('role', 'role').replace(' ', '_')}_job_description_{date.today()}.docx"
        return self._save_docx(doc, filename)

    def _fill_template(self, template_path: str, fields: dict) -> str:
        """
        Replace {{PLACEHOLDER}} tokens in a .docx template.
        Returns the output file path.
        """
        print(f"Looking for template at: {template_path}")
        print(f"File exists: {os.path.exists(template_path)}")
        doc = Document(template_path)
        normalized = {k.upper(): str(v) for k, v in fields.items()}

        # Add today's date if not provided
        if "DATE" not in normalized:
            normalized["DATE"] = date.today().strftime("%B %d, %Y")

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, normalized)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, normalized)

        return doc

    def _replace_in_paragraph(self, paragraph, fields: dict):
        """Replace placeholders in a paragraph while preserving formatting."""
        full_text = "".join(run.text for run in paragraph.runs)
    
        replaced = full_text
        for key, value in fields.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in replaced:
                replaced = replaced.replace(placeholder, value)
    
        if replaced == full_text:
            return  # nothing changed, leave runs untouched
    
        # Collapse double spaces that can result from an empty-string field
        # value sitting between two other words (see note above).
        replaced = re.sub(r" {2,}", " ", replaced).strip()
    
        if not paragraph.runs:
            return
    
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""

    def _save_docx(self, doc, filename: str) -> str:
        """Save a filled document to the output directory."""
        output_path = os.path.join(self.output_dir, filename)
        doc.save(output_path)
        return output_path

    def get_placeholder_contexts(self, template_path: str) -> dict:
        """
        Map each {{PLACEHOLDER}} key in the template to the literal sentence/
        paragraph text it appears in. Used by DraftingAgent so the model sees
        each blank in its actual sentence context, rather than just a bare
        field name — this is what keeps it from drafting a full restated
        clause when only a short fragment is needed (e.g. "on a biweekly
        basis" instead of "Payment will be made on a biweekly basis.").
    
        If the same key appears in multiple paragraphs (e.g. {{COMPANY_NAME}}
        showing up in both the letterhead and the signature block), the first
        occurrence's context is used — that's sufficient for drafting purposes
        since the value itself is identical wherever the key repeats.
        """
        doc = Document(template_path)
        contexts = {}
    
        def scan_paragraphs(paragraphs):
            for paragraph in paragraphs:
                keys = re.findall(r"\{\{([A-Z_]+)\}\}", paragraph.text)
                for key in keys:
                    if key not in contexts:
                        contexts[key] = paragraph.text
    
        scan_paragraphs(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    scan_paragraphs(cell.paragraphs)
    
        return contexts
    
    
    def render_text_preview(self, template_path: str, fields: dict) -> str:
        """
        Fill a template's placeholders the same way _fill_template does, then
        return the result as plain text (paragraphs joined by newlines, table
        cells included) instead of saving a .docx. Used to render the UI draft
        preview from the exact same field values used for document generation.
        """
        doc = Document(template_path)
        normalized = {k.upper(): str(v) for k, v in fields.items()}
    
        if "DATE" not in normalized:
            normalized["DATE"] = date.today().strftime("%B %d, %Y")
    
        lines = []
    
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, normalized)
            lines.append(paragraph.text)
    
        for table in doc.tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, normalized)
                    row_cells.append(cell.text)
                lines.append(" | ".join(row_cells))
    
        return "\n".join(lines)

    def list_unfilled_placeholders(self, file_path: str) -> list[str]:
        """Scan a .docx template and return all placeholder keys."""
        doc = Document(file_path)
        placeholders = []
        full_text = "\n".join([p.text for p in doc.paragraphs])
        return re.findall(r"\{\{([A-Z_]+)\}\}", full_text)
    